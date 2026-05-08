"""
docx_exporter.py
================
Converts a PageNode document tree into a DOCX file (bytes).
Requires python-docx:  pip install python-docx

LaTeX formulas are rendered as images via matplotlib if available,
otherwise the raw LaTeX string is inserted as monospace text.
"""

from __future__ import annotations

import io
from pathlib import Path
from typing import Optional

from schema import PageNode, ContentGroupNode, Block


def _latex_to_png(latex: str) -> Optional[bytes]:
    """Render a LaTeX expression to a PNG image using matplotlib."""
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        fig = plt.figure(figsize=(6, 0.8))
        fig.text(0.5, 0.5, f"${latex.strip()}$", ha="center", va="center", fontsize=16)
        buf = io.BytesIO()
        fig.savefig(buf, format="png", bbox_inches="tight", dpi=150,
                    facecolor="white", transparent=False)
        plt.close(fig)
        buf.seek(0)
        return buf.read()
    except Exception:
        return None


def _add_block_to_para(para, block: Block, doc, img_dir: Path) -> None:
    """Add a single content block's content to a paragraph (or document)."""
    from docx.shared import Pt, Inches
    from docx.oxml.ns import qn
    import docx

    t = block.type
    if t == "text":
        para.add_run(block.value + " ")
    elif t == "latex":
        png = _latex_to_png(block.value)
        if png:
            run = para.add_run()
            run.add_picture(io.BytesIO(png), width=Inches(3.5 if block.display else 1.5))
        else:
            run = para.add_run(f"{'$$' if block.display else '$'}{block.value}{'$$' if block.display else '$'} ")
            run.font.name = "Courier New"
            run.font.size = Pt(10)
    elif t == "image":
        # Try to resolve local image path
        url = block.url
        img_path: Optional[Path] = None
        if url.startswith("/static/"):
            candidate = img_dir.parent / url.lstrip("/").replace("/", "/")
            if candidate.exists():
                img_path = candidate
        if img_path:
            run = para.add_run()
            try:
                run.add_picture(str(img_path), width=Inches(4))
            except Exception:
                para.add_run(f"[image: {url}]")
        else:
            para.add_run(f"[image: {url}]")


def _add_group(doc, grp: ContentGroupNode, img_dir: Path) -> None:
    from docx.shared import Pt
    import docx

    if grp.label == "title":
        heading = doc.add_heading("", level=2)
        for b in grp.blocks:
            if b.type == "text":
                heading.add_run(b.value + " ")
        return

    if grp.label == "table" and grp.blocks:
        b = grp.blocks[0]
        if b.type == "table" and b.cells:
            cols = max((len(row) for row in b.cells), default=1)
            tbl = doc.add_table(rows=len(b.cells), cols=cols)
            tbl.style = "Table Grid"
            for ri, row in enumerate(b.cells):
                for ci, cell_blocks in enumerate(row):
                    cell_text = " ".join(
                        cb.value for cb in cell_blocks if cb.type == "text"
                    )
                    tbl.rows[ri].cells[ci].text = cell_text
            doc.add_paragraph()
            return

    para = doc.add_paragraph()
    for b in grp.blocks:
        _add_block_to_para(para, b, doc, img_dir)


def export_docx(page: PageNode, img_dir: Path = Path(".")) -> bytes:
    try:
        import docx
    except ImportError:
        raise ImportError("python-docx is required: pip install python-docx")

    doc = docx.Document()

    for col in page.columns:
        if len(page.columns) > 1:
            doc.add_heading(f"Column {col.column_idx}", level=3)
        for grp in col.groups:
            _add_group(doc, grp, img_dir)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
